# export_utils.py
from __future__ import annotations

import base64
import html
import io
import mimetypes
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

# Emplacement des fichiers upload (images de texte)
try:
    from paths import UPLOAD_DIR
except Exception:
    UPLOAD_DIR = Path(os.getcwd()) / "uploads"


# ============== Helpers ==============

def _escape(s: Optional[str]) -> str:
    return html.escape(s or "")


def _nl2br_escaped(s: Optional[str]) -> str:
    return _escape(s).replace("\n", "<br>")


def _read_bytes(p: Path) -> Optional[bytes]:
    try:
        return p.read_bytes()
    except Exception:
        return None


def _to_data_uri_from_bytes(data: bytes, mime: str) -> str:
    return f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"


def _to_data_uri_from_path(p: Path) -> Optional[str]:
    b = _read_bytes(p)
    if not b:
        return None
    mime = mimetypes.guess_type(str(p))[0] or "application/octet-stream"
    return _to_data_uri_from_bytes(b, mime)


def _absolute_url(url: Optional[str], base_url: Optional[str]) -> Optional[str]:
    """
    - http(s) → renvoyé tel quel
    - "/uploads/..." + base_url → base_url + path
    - sinon valeur brute
    """
    if not url:
        return None
    u = str(url)
    if u.startswith(("http://", "https://")):
        return u
    if base_url and u.startswith("/"):
        return base_url.rstrip("/") + u
    return u


def _human_dt(iso: Optional[str]) -> str:
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
    except Exception:
        return iso
    return dt.strftime("%d/%m/%Y %H:%M")


def _logo_data_uri(logo_path: Optional[Path]) -> str:
    """
    Retourne un data:URI base64 pour le logo.
    Utilise logo_path si fourni/valide ; sinon fallback PNG 1x1 (sûr pour <img src="...">).
    """
    if logo_path and logo_path.exists():
        data = logo_path.read_bytes()
        mime = mimetypes.guess_type(str(logo_path))[0] or "image/png"
        return _to_data_uri_from_bytes(data, mime)

    # fallback PNG transparent 1x1
    _PNG_1x1 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO9Z9yQAAAAASUVORK5CYII="
    return "data:image/png;base64," + _PNG_1x1


# ============== HTML (fidèle à ton test) ==============

# IMPORTANT : on utilise %-format → doubler les % dans le CSS (%%)
_HTML_FLEX = """\
<!doctype html>
<html lang="fr">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>PastelNotes — Export</title>
<style>
  :root{
    --side-w: 120px;
    --page-bg: #ffffff;
    --pastel-left: linear-gradient(180deg,#E9FCE0 0%%, #D9F8E4 100%%);
    --topbar-gray: #EDEFF2;
    --shadow: 0 10px 30px rgba(16,32,64,0.06);
    --muted: #7a8590;
  }
  html,body{height:100%%; margin:0; font-family: "Inter","Helvetica Neue", Arial, sans-serif; color:#111; background:#f6f8f9;}
  .paper {
    /* A4 visual size + shadow (écran) */
    width: 210mm;
    min-height: 297mm;
    margin: 0;                    /* pas de marge écran (PDF marges=0) */
    background: var(--page-bg);
    box-shadow: none;             /* désactivé pour le PDF plein format */
    box-sizing: border-box;
    display: flex;                /* <-- flex layout */
    flex-direction: row;
    border-radius: 0;
    overflow: hidden;
  }

  /* LEFT SIDEBAR */
  .side {
    width: var(--side-w);
    min-width: var(--side-w);
    background: var(--pastel-left);
    display: flex;
    flex-direction: column;   /* stack vertically */
    align-items: flex-start;  /* logo + texts aligned to left */
    justify-content: flex-start;
    padding: 20px 12px;
    box-sizing: border-box;
  }

  /* Make the left column visually span whole page height */
  .side .brand {
    display:flex;
    flex-direction: column;
    gap:8px;
    align-items: flex-start;
    width: 100%%;
    margin: auto;
  }
  .logo {
    width:100%%; height:100%%;
    border-radius:8px;
    overflow:hidden;
    display:block;
  }
  .logo img{ width:100%%; height:100%%; object-fit:contain; display:block; }

  .brand-name{
    font-weight:700;
    color:#2d8aa8;
    font-size:14px;
    line-height:1;
    margin:0;
  }
  .brand-sub{
    font-size:11px;
    color:#3f6b6b;
    margin-top:2px;
  }

  /* push footer info to bottom */
  .side .spacer{ flex:1 1 auto; } /* takes remaining vertical space */
  .side .meta-bottom{
    font-size:11px; color:#9aa6a6; margin-top:6px;
  }

  /* MAIN CONTENT */
  .main {
    flex: 1 1 auto;
    padding: 18px 28px;
    box-sizing: border-box;
    display: flex;
    flex-direction: column;
  }

  .topbar{
    width:100%%;
    height:120px;
    border-radius:8px;
    overflow:hidden;
    background: var(--topbar-gray);
    margin-bottom:18px;
    background-size:cover;
    background-position:center;
  }

  .content{
    flex: 1 1 auto;
  }
  h1{font-size:34px; margin:4px 0 8px 0; font-weight:700;}
  .meta{color:var(--muted); font-size:13px; margin-bottom:14px;}
  .context{font-style:italic; color:#7a8590; margin:10px 0;}
  .body{margin-top:12px; line-height:1.6; font-size:15px; white-space:pre-wrap;}

  .footer{
    display:flex; justify-content:space-between; margin-top:28px; color:var(--muted); font-size:12px; border-top:1px solid #f2f3f4; padding-top:10px;
  }

  /* Impression plein format avec Chromium */
  @media print {
    html,body{ margin:0; }
    .paper{
      margin:0;
      width:210mm; height:297mm;     /* prend TOUTE la hauteur */
      border-radius:0; box-shadow:none;
    }
  }
  @page { size: A4; margin: 0; }     /* Aucune marge PDF */
</style>
</head>
<body>
  <div class="paper">
    <div class="side" role="complementary" aria-label="colonne">
      <div class="brand">
        <a class="logo" href="#" aria-hidden="true">
          <img src="%(logo_src)s" alt="logo PastelNotes">
        </a>
        <div class="brand-name">PastelNotes</div>
        <div class="brand-sub">Export élégant</div>
      </div>

      <div class="spacer" aria-hidden="true"></div>

      <div class="meta-bottom">
        Export PastelNotes<br>
        %(stamp)s
      </div>
    </div>

    <main class="main" role="main">
      <div class="topbar" style="%(topbar_style)s"></div>

      <section class="content" aria-labelledby="title">
        <h1 id="title">%(title)s</h1>
        <div class="meta">Par <strong>%(author)s</strong> — %(date_h)s</div>

        <div class="block">
          <span style="display:inline-block; width:26px; height:26px; line-height:26px; text-align:center; border-radius:50%%; background:#f0f8ff; margin-right:8px; color:#2d8aa8; font-weight:700;">♪</span>
          <strong>Source :</strong>
          %(music_link_html)s
        </div>

        %(context_html)s

        <div class="body">
%(body_html)s
        </div>
      </section>

      <div class="footer">
        <div>PastelNotes — Export</div>
        <div>Par %(author)s — %(date_h)s</div>
      </div>
    </main>
  </div>
</body>
</html>
"""


def build_export_html(
    text: Dict[str, Any],
    *,
    logo_path: Optional[Path] = None,
    base_url: Optional[str] = None,
) -> str:
    """
    Construit le HTML final *identique* à ton HTML de test (flex + gradient + cover 150%%),
    prêt à être imprimé par Chromium headless (Playwright).
    """
    # Logo (data-URI)
    logo_src = _logo_data_uri(logo_path)

    # Cover : image upload → data URI ; URL distante → http(s)
    topbar_style = "background-color:#EDEFF2;"
    cover_uri: Optional[str] = None
    if text.get("image_filename"):
        p = UPLOAD_DIR / str(text["image_filename"])
        cover_uri = _to_data_uri_from_path(p)
    elif text.get("image_url"):
        cover_uri = _absolute_url(text["image_url"], base_url)

    if cover_uri:
        # explicit props so shorthand doesn’t reset them
        topbar_style = (
            f"background-color:#EDEFF2;"
            f"background-image:url('{cover_uri}');"
            "background-position:center center;"
            "background-size:cover;"
            "background-repeat:no-repeat;"
        )

    # Musique
    music = text.get("music_original_url") or text.get("music_url") or ""
    music_abs = _absolute_url(music, base_url)
    music_link_html = (
        f'<a href="{_escape(music_abs)}" target="_blank">{_escape(music_abs)}</a>'
        if music_abs else "<span style='color:#7a8590'>—</span>"
    )

    # Contexte / Corps
    context_html = f"<div class='context'>Contexte : {_nl2br_escaped(text.get('context'))}</div>" if text.get("context") else ""
    body_html = _nl2br_escaped(text.get("body") or "")

    mapping = {
        "logo_src": logo_src,
        "stamp": datetime.utcnow().strftime("%Y-%m-%d_%H:%M"),
        "topbar_style": topbar_style,  # <-- add this
        "title": _escape(text.get("title") or "(sans titre)"),
        "author": _escape(text.get("created_by") or "—"),
        "date_h": _human_dt(text.get("date")),
        "music_link_html": music_link_html,
        "context_html": context_html,
        "body_html": body_html,
    }
    return _HTML_FLEX % mapping


# ============== Rendu PDF : Chromium/Playwright (fallbacks) ==============

def _render_pdf_with_playwright(html_str: str) -> bytes:
    """Rendu PDF via Chromium headless (Playwright) SANS marges, fond imprimé, media print."""
    from playwright.sync_api import sync_playwright  # type: ignore
    print("[export][pdf] engine=playwright")  # DEBUG

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True, args=["--no-sandbox", "--disable-dev-shm-usage"])
        context = browser.new_context(no_viewport=True, ignore_https_errors=True)
        page = context.new_page()
        # page.set_cache_enabled(False)  # 🔥 empêche tout cache HTML/CSS
        page.emulate_media(media="print")  # applique @media print et @page
        print("[export] regenerate CSS (timestamp):", datetime.utcnow())
        page.set_content(html_str, wait_until="networkidle")
        pdf = page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},  # AUCUNE marge
            prefer_css_page_size=True,
            scale=1.0,
        )
        browser.close()
        return pdf


def export_pdf_bytes(
    text: Dict[str, Any],
    *,
    logo_path: Optional[Path] = None,
    base_url: Optional[str] = None,
) -> bytes:
    """
    Ordre de rendu :
      1) Playwright/Chromium (fidèle, sans marges)
      2) WeasyPrint
      3) wkhtmltopdf/pdfkit
    → imprime en debug le moteur utilisé.
    """
    html_str = build_export_html(text, logo_path=logo_path, base_url=base_url)
    errors = []

    # 1) Chromium
    try:
        return _render_pdf_with_playwright(html_str)
    except Exception as e:
        errors.append(f"Playwright: {e}")

    # 2) WeasyPrint
    try:
        from weasyprint import HTML  # type: ignore
        print("[export][pdf] engine=weasyprint")  # DEBUG
        return HTML(string=html_str, base_url=base_url or ".").write_pdf()
    except Exception as e:
        errors.append(f"WeasyPrint: {e}")

    # 3) wkhtmltopdf/pdfkit
    try:
        import pdfkit  # type: ignore
        print("[export][pdf] engine=wkhtmltopdf")  # DEBUG
        opts = {
            "quiet": "",
            "enable-local-file-access": "",
            "encoding": "UTF-8",
            "print-media-type": "",
            "page-size": "A4",
            "margin-top": "0",
            "margin-right": "0",
            "margin-bottom": "0",
            "margin-left": "0",
        }
        return pdfkit.from_string(html_str, False, options=opts)
    except Exception as e:
        errors.append(f"pdfkit: {e}")

    raise RuntimeError("Impossible de générer le PDF. " + " | ".join(errors))


# ============== DOCX (sobre) ==============

def export_docx_bytes(
    text: Dict[str, Any],
    *,
    logo_path: Optional[Path] = None
) -> bytes:
    """
    DOCX sobre ; le PDF Playwright reste la cible “pixel-perfect”.
    """
    from docx import Document  # type: ignore
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # header avec logo si fourni
    if logo_path and logo_path.exists():
        try:
            hdr = doc.sections[0].header
            p = hdr.paragraphs[0]
            run = p.add_run()
            run.add_picture(str(logo_path), width=Inches(0.7))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    title = text.get("title") or "(sans titre)"
    author = text.get("created_by") or "—"
    date_h = _human_dt(text.get("date"))

    p = doc.add_paragraph(title)
    p.style = doc.styles["Title"]

    meta = doc.add_paragraph(f"Par {author} — {date_h}")
    meta.runs[0].font.size = Pt(10)

    link = text.get("music_original_url") or text.get("music_url")
    if link:
        doc.add_paragraph(f"Source : {link}")

    if text.get("context"):
        para = doc.add_paragraph()
        r = para.add_run(f"Contexte : {text['context']}")
        r.italic = True

    # image locale si dispo
    if text.get("image_filename"):
        pth = UPLOAD_DIR / str(text["image_filename"])
        if pth.exists():
            try: doc.add_picture(str(pth), width=Inches(5.5))
            except Exception: pass

    for line in str(text.get("body") or "").splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
